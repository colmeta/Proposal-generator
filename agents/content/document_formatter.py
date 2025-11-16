"""
Document Formatter Agent
Professional formatting, document structure, export to PDF/DOCX, style consistency
"""

from typing import Dict, Any, List, Optional
from agents.base_agent import BaseAgent
from config.llm_config import LLMProvider
import json
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False


class DocumentFormatterAgent(BaseAgent):
    """
    Document Formatter Agent
    Handles professional formatting and export to PDF/DOCX
    """
    
    def __init__(self):
        super().__init__(
            name="Document Formatter",
            role="Professional document formatting and export",
            task_type="general"
        )
        self.default_style = {
            "font_family": "Arial",
            "font_size": 11,
            "line_spacing": 1.15,
            "margins": {"top": 1, "bottom": 1, "left": 1, "right": 1},
            "heading_styles": {
                "h1": {"size": 18, "bold": True, "color": (0, 0, 0)},
                "h2": {"size": 16, "bold": True, "color": (0, 0, 0)},
                "h3": {"size": 14, "bold": True, "color": (0, 0, 0)}
            }
        }
    
    def format_document(
        self,
        content: Dict[str, Any],
        style_guide: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Format document with professional styling
        
        Args:
            content: Document content with sections
            style_guide: Custom style guide (optional)
        
        Returns:
            Formatted document structure
        """
        self.log_action("Formatting document")
        
        style = style_guide if style_guide else self.default_style
        
        # Structure the document
        formatted = {
            "title": content.get("title", "Document"),
            "sections": [],
            "style": style,
            "metadata": content.get("metadata", {})
        }
        
        sections = content.get("sections", {})
        if isinstance(sections, dict):
            for section_name, section_content in sections.items():
                if isinstance(section_content, dict):
                    section_text = section_content.get("content", "")
                else:
                    section_text = str(section_content)
                
                formatted["sections"].append({
                    "name": section_name,
                    "content": section_text,
                    "level": 1,
                    "formatted": True
                })
        
        self.log_action("Document formatted", {"sections": len(formatted["sections"])})
        return formatted
    
    def export_to_docx(
        self,
        document: Dict[str, Any],
        output_path: str
    ) -> Dict[str, Any]:
        """
        Export document to DOCX format
        
        Args:
            document: Formatted document structure
            output_path: Path to save DOCX file
        
        Returns:
            Export result with file path
        """
        self.log_action(f"Exporting to DOCX: {output_path}")
        
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "python-docx not available. Install with: pip install python-docx",
                "output_path": output_path
            }
        
        try:
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(document.get("style", {}).get("margins", {}).get("top", 1))
                section.bottom_margin = Inches(document.get("style", {}).get("margins", {}).get("bottom", 1))
                section.left_margin = Inches(document.get("style", {}).get("margins", {}).get("left", 1))
                section.right_margin = Inches(document.get("style", {}).get("margins", {}).get("right", 1))
            
            # Add title
            title = doc.add_heading(document.get("title", "Document"), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add sections
            for section in document.get("sections", []):
                section_name = section.get("name", "")
                section_content = section.get("content", "")
                
                # Add section heading
                if section_name:
                    heading = doc.add_heading(section_name, level=1)
                
                # Add section content
                if section_content:
                    # Split into paragraphs
                    paragraphs = section_content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            para = doc.add_paragraph(para_text.strip())
                            para_format = para.paragraph_format
                            para_format.line_spacing = document.get("style", {}).get("line_spacing", 1.15)
                    
                    # Add spacing after section
                    doc.add_paragraph()
            
            # Ensure output directory exists
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
            
            # Save document
            doc.save(output_path)
            
            self.log_action("DOCX exported successfully", {"path": output_path})
            return {
                "success": True,
                "output_path": output_path,
                "format": "docx",
                "sections": len(document.get("sections", []))
            }
        except Exception as e:
            self.logger.error(f"DOCX export failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "output_path": output_path
            }
    
    def export_to_pdf(
        self,
        document: Dict[str, Any],
        output_path: str
    ) -> Dict[str, Any]:
        """
        Export document to PDF format
        
        Args:
            document: Formatted document structure
            output_path: Path to save PDF file
        
        Returns:
            Export result with file path
        """
        self.log_action(f"Exporting to PDF: {output_path}")
        
        if not REPORTLAB_AVAILABLE:
            return {
                "success": False,
                "error": "reportlab not available. Install with: pip install reportlab",
                "output_path": output_path
            }
        
        try:
            # Create PDF document
            pdf_doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=inch * document.get("style", {}).get("margins", {}).get("right", 1),
                leftMargin=inch * document.get("style", {}).get("margins", {}).get("left", 1),
                topMargin=inch * document.get("style", {}).get("margins", {}).get("top", 1),
                bottomMargin=inch * document.get("style", {}).get("margins", {}).get("bottom", 1)
            )
            
            # Build content
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=RGBColor(0, 0, 0),
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                textColor=RGBColor(0, 0, 0),
                spaceAfter=12,
                spaceBefore=12
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=11,
                leading=13,
                alignment=TA_JUSTIFY
            )
            
            # Add title
            title = document.get("title", "Document")
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.3*inch))
            
            # Add sections
            for section in document.get("sections", []):
                section_name = section.get("name", "")
                section_content = section.get("content", "")
                
                # Add section heading
                if section_name:
                    story.append(Paragraph(section_name, heading_style))
                    story.append(Spacer(1, 0.1*inch))
                
                # Add section content
                if section_content:
                    # Split into paragraphs
                    paragraphs = section_content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            # Escape special characters for ReportLab
                            para_text_escaped = para_text.strip().replace('&', '&amp;')
                            story.append(Paragraph(para_text_escaped, body_style))
                            story.append(Spacer(1, 0.1*inch))
                    
                    story.append(Spacer(1, 0.2*inch))
            
            # Ensure output directory exists
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
            
            # Build PDF
            pdf_doc.build(story)
            
            self.log_action("PDF exported successfully", {"path": output_path})
            return {
                "success": True,
                "output_path": output_path,
                "format": "pdf",
                "sections": len(document.get("sections", []))
            }
        except Exception as e:
            self.logger.error(f"PDF export failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "output_path": output_path
            }
    
    def ensure_style_consistency(
        self,
        document: Dict[str, Any],
        style_guide: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Ensure style consistency throughout document
        
        Args:
            document: Document to check
            style_guide: Style guide to apply
        
        Returns:
            Consistency check results
        """
        self.log_action("Checking style consistency")
        
        issues = []
        sections_checked = 0
        
        for section in document.get("sections", []):
            sections_checked += 1
            # Check if section follows style guide
            # This is a simplified check - could be enhanced
            if not section.get("formatted", False):
                issues.append(f"Section '{section.get('name')}' not properly formatted")
        
        consistency_score = 10.0 - (len(issues) * 2) if issues else 10.0
        consistency_score = max(0, consistency_score)
        
        return {
            "consistent": len(issues) == 0,
            "consistency_score": consistency_score,
            "issues": issues,
            "sections_checked": sections_checked,
            "recommendations": ["Apply style guide consistently"] if issues else []
        }
    
    def create_document_structure(
        self,
        sections: Dict[str, Any],
        title: str = "Document"
    ) -> Dict[str, Any]:
        """
        Create proper document structure
        
        Args:
            sections: Document sections
            title: Document title
        
        Returns:
            Structured document
        """
        self.log_action("Creating document structure")
        
        structured = {
            "title": title,
            "sections": [],
            "style": self.default_style,
            "metadata": {
                "created_by": "Document Formatter Agent",
                "format": "structured"
            }
        }
        
        for section_name, section_content in sections.items():
            if isinstance(section_content, dict):
                content = section_content.get("content", "")
            else:
                content = str(section_content)
            
            structured["sections"].append({
                "name": section_name,
                "content": content,
                "level": 1,
                "formatted": False
            })
        
        return structured
    
    def process(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """Process input - route to appropriate method"""
        action = input_data.get("action", "format")
        
        if action == "format":
            return self.format_document(
                input_data.get("content", {}),
                input_data.get("style_guide")
            )
        elif action == "export_docx":
            return self.export_to_docx(
                input_data.get("document", {}),
                input_data.get("output_path", "output.docx")
            )
        elif action == "export_pdf":
            return self.export_to_pdf(
                input_data.get("document", {}),
                input_data.get("output_path", "output.pdf")
            )
        elif action == "check_consistency":
            return self.ensure_style_consistency(
                input_data.get("document", {}),
                input_data.get("style_guide", self.default_style)
            )
        elif action == "create_structure":
            return self.create_document_structure(
                input_data.get("sections", {}),
                input_data.get("title", "Document")
            )
        else:
            # Default to formatting
            return self.format_document(
                input_data.get("content", {}),
                input_data.get("style_guide")
            )

